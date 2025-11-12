#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown转DOCX脚本
用于将Markdown文件转换为DOCX格式的文档

使用方法:
    python md_to_docx.py <markdown_file_path> [output_dir]

示例:
    python md_to_docx.py ../public/pdfs/ethics/Assignment\\ 2/Assignment\\ 2\\ Response.md
"""

from __future__ import annotations

import argparse
import re
import sys
from io import BytesIO
from pathlib import Path
from typing import Optional

from html2docx import html2docx
from markdown import markdown
from docx import Document


class MarkdownToDocxConverter:
    """Markdown转DOCX转换器"""

    def __init__(self) -> None:
        self.supported_formats = {".md", ".markdown"}

    def read_markdown(self, markdown_path: Path) -> str:
        """
        读取Markdown文件内容

        Args:
            markdown_path: Markdown文件路径

        Returns:
            str: Markdown文本
        """
        try:
            return markdown_path.read_text(encoding="utf-8")
        except Exception as exc:  # pylint: disable=broad-except
            raise ValueError(f"无法读取Markdown文件 {markdown_path}: {exc}") from exc

    def convert_markdown_to_docx(
        self,
        markdown_path: str,
        output_dir: Optional[str] = None,
    ) -> str:
        """
        将Markdown转换为DOCX文件

        Args:
            markdown_path: Markdown文件路径
            output_dir: 输出目录（可选）

        Returns:
            str: 生成的DOCX文件路径
        """
        src_path = Path(markdown_path)

        if not src_path.exists():
            raise FileNotFoundError(f"Markdown文件不存在: {src_path}")

        if src_path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {src_path.suffix}")

        if output_dir:
            dst_dir = Path(output_dir)
            dst_dir.mkdir(parents=True, exist_ok=True)
        else:
            dst_dir = src_path.parent

        output_path = dst_dir / f"{src_path.stem}.docx"

        print(f"正在处理Markdown文件: {src_path}")
        print(f"输出文件: {output_path}")

        markdown_text = self.read_markdown(src_path)
        if not markdown_text.strip():
            raise ValueError("Markdown文件内容为空")

        html = markdown(
            markdown_text,
            extensions=["extra", "tables", "fenced_code", "toc"],
        )
        html = self.clean_html(html)
        docx_buffer = html2docx(html, src_path.stem)
        docx_buffer = self.post_process_docx(docx_buffer, markdown_text)
        with output_path.open("wb") as file:
            file.write(docx_buffer.getvalue())

        print(f"DOCX文件已生成: {output_path}")
        return str(output_path)

    @staticmethod
    def clean_html(html: str) -> str:
        """清理转换后的HTML，移除空段落等冗余内容"""
        empty_paragraph_pattern = re.compile(
            r"<p[^>]*>(?:\s|&nbsp;|<br\s*/?>)*</p>", flags=re.IGNORECASE
        )
        return re.sub(empty_paragraph_pattern, "", html)

    @staticmethod
    def post_process_docx(docx_buffer: BytesIO, markdown_text: str) -> BytesIO:
        """
        对生成的DOCX进行后处理：移除空白段落、修正标题样式

        Args:
            docx_buffer: DOCX内容的内存缓冲区
            markdown_text: 原始Markdown文本

        Returns:
            BytesIO: 已清理的DOCX缓冲区
        """
        docx_buffer.seek(0)
        document = Document(docx_buffer)
        modified = False

        if MarkdownToDocxConverter.apply_heading_styles(document, markdown_text):
            modified = True

        for paragraph in list(document.paragraphs):
            if not paragraph.text.strip():
                paragraph_element = paragraph._element  # pylint: disable=protected-access
                parent_element = paragraph_element.getparent()
                parent_element.remove(paragraph_element)
                modified = True

        if not modified:
            docx_buffer.seek(0)
            return docx_buffer

        cleaned_buffer = BytesIO()
        document.save(cleaned_buffer)
        cleaned_buffer.seek(0)
        return cleaned_buffer

    @staticmethod
    def apply_heading_styles(document: Document, markdown_text: str) -> bool:
        """根据Markdown内容修正DOCX中的标题样式"""
        headings = MarkdownToDocxConverter.extract_headings(markdown_text)
        if not headings:
            return False

        modified = False
        heading_index = 0
        for paragraph in document.paragraphs:
            if heading_index >= len(headings):
                break

            level, heading_text = headings[heading_index]
            paragraph_text = paragraph.text.strip()
            if paragraph_text != heading_text:
                continue

            target_style_name = f"Heading {level}"
            if paragraph.style.name != target_style_name:
                try:
                    paragraph.style = document.styles[target_style_name]
                    modified = True
                except KeyError:
                    # 如果目标样式不存在，则跳过但继续匹配后续标题
                    pass
            heading_index += 1

        return modified

    @staticmethod
    def extract_headings(markdown_text: str) -> list[tuple[int, str]]:
        """提取Markdown中的标题层级和文本"""
        heading_pattern = re.compile(r"^(#{1,6})\s+(.*)$")
        headings: list[tuple[int, str]] = []
        for line in markdown_text.splitlines():
            match = heading_pattern.match(line)
            if not match:
                continue
            level = len(match.group(1))
            text = match.group(2).strip()
            text = re.sub(r"\s+#*$", "", text)
            if text:
                headings.append((level, text))

        return headings


def main() -> None:
    """主函数"""
    parser = argparse.ArgumentParser(
        description="将Markdown文件转换为DOCX格式",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例用法:
  python md_to_docx.py ../public/pdfs/ethics/Assignment\\ 2/Assignment\\ 2\\ Response.md
  python md_to_docx.py ../public/pdfs/ethics/Assignment\\ 2/Assignment\\ 2\\ Response.md ./output
        """,
    )

    parser.add_argument(
        "markdown_path",
        help="Markdown文件路径",
    )

    parser.add_argument(
        "output_dir",
        nargs="?",
        help="输出目录（可选，默认为Markdown文件所在目录）",
    )

    args = parser.parse_args()

    converter = MarkdownToDocxConverter()

    try:
        src_path = Path(args.markdown_path)
        if src_path.is_file():
            converter.convert_markdown_to_docx(str(src_path), args.output_dir)
        else:
            print(f"错误: 文件不存在或不是有效的Markdown文件: {src_path}")
            sys.exit(1)
    except Exception as exc:  # pylint: disable=broad-except
        print(f"错误: {exc}")
        sys.exit(1)


if __name__ == "__main__":
    main()

